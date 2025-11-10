import os
import sys
from docx import Document

# Set UTF-8 encoding for output
sys.stdout.reconfigure(encoding='utf-8')

# Get the current directory
current_dir = os.path.dirname(os.path.abspath(__file__))

# Find the .docx file with proper encoding
try:
    files = os.listdir(current_dir)
    docx_files = [f for f in files if f.endswith('.docx')]
    
    if docx_files:
        # Try to open the file
        for docx_file in docx_files:
            try:
                file_path = os.path.join(current_dir, docx_file)
                print(f"Trying to open: {docx_file}", file=sys.stderr)
                doc = Document(file_path)
                print("\n--- DOCUMENT CONTENT ---\n")
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        print(text)
                
                # Also check tables
                if doc.tables:
                    print("\n--- TABLES ---\n")
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                print(row_text)
                break
            except Exception as e:
                print(f"Error opening {docx_file}: {e}", file=sys.stderr)
                continue
    else:
        print("No .docx file found in directory", file=sys.stderr)
        print(f"Files in directory: {files}", file=sys.stderr)
except Exception as e:
    print(f"Error: {e}", file=sys.stderr)

